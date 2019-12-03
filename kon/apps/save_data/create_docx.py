from docxtpl import DocxTemplate
from docx import Document
import os
from docxtpl import DocxTemplate
from io import StringIO


def cr_docx():
    path = (os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    print(path)

    # f = open(os.path.join(path, 'report_template.docx'), 'rb')
    doc = DocxTemplate('../../templates/report_template.docx')
    context = {'f': "k_sub"}
    doc.render(context)
    doc.save("../../templates/report22222.docx")
    # f.close()


    # f = open(os.path.join(path, 'report_template.docx'), 'rb')
    # doc = Document(f)
    # p = doc.add_paragraph('Объект госавиации: ' + 'test worlds!!!!!!')
    # # f1 = StringIO()
    # doc.save('../test_doc/report2.docx')
    # f.close()

    # with open(os.path.join(path, 'report_template.docx'), 'r') as f:
    #     source_stream = StringIO(f.read())
    # doc = Document(source_stream)
    # p = doc.add_paragraph('Объект госавиации: ' + 'test worlds!!!!!!')
    # source_stream.close()
    # target_stream = StringIO()
    # doc.save(target_stream)



    # doc = Document()
    # p = doc.add_paragraph('Объект госавиации: ' + 'test worlds!!!!!!')
    # doc.save('report_template.docx')


cr_docx()


# print((os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

# path = ((os.path.dirname(os.path.abspath(__file__))))
# print(path)



